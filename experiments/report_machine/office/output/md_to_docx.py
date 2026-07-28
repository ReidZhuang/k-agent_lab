#!/usr/bin/env python3
"""
Markdown → DOCX 转换工具（VS Code Markdown Preview 风格）

模仿 VS Code 内置 Markdown Preview 的干净、现代排版风格。

特性:
  - H1/H2 带底部通栏细线（仿 GitHub/VS Code 风格）
  - 字体：无衬线体，标题深灰 #1a1a1a，正文 #333
  - 表格：最小化边框，表头浅灰底 #f6f8fa
  - 封面页简洁克制
  - 中文排版优化

用法:
  conda run -n stock_agent python md_to_docx.py input.md output.docx
  conda run -n stock_agent python md_to_docx.py --dir output/
"""
import os
import sys
import argparse
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── VS Code Markdown Preview 风格配色 ──
C_H1 = RGBColor(0x1A, 0x1A, 0x1A)          # H1: 最深灰
C_H2 = RGBColor(0x1A, 0x1A, 0x1A)          # H2: 同 H1（字号区分）
C_H3 = RGBColor(0x33, 0x33, 0x33)          # H3: 深灰
C_BODY = RGBColor(0x33, 0x33, 0x33)        # 正文: #333
C_MUTED = RGBColor(0x66, 0x66, 0x66)        # 辅助: #666
C_LIGHT = RGBColor(0x99, 0x99, 0x99)        # 更浅
C_LINK = RGBColor(0x03, 0x6A, 0xD7)         # 链接蓝
C_TABLE_BORDER = "D0D7DE"                   # 表格边框
C_TABLE_HEADER = "F6F8FA"                   # 表头背景
C_TABLE_ALT = "F8F9FA"                      # 交替行
C_DIVIDER = "D0D7DE"                        # 分割线颜色


def md_to_docx(md_text: str, stock_name: str = "",
               trade_date: str = "") -> Document:
    """将 Markdown 报告文本转换为 Word 文档

    Args:
        md_text: 报告的 Markdown 内容
        stock_name: 股票名称（用于封面）
        trade_date: 交易日期（用于封面）

    Returns:
        python-docx Document 实例
    """
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # ── 页眉页脚 ──
    _add_header_footer(doc, stock_name)

    # ── 默认样式 ──
    _set_styles(doc)

    # ── 封面页 ──
    _add_cover(doc, stock_name, trade_date, md_text)

    # ── 正文 ──
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # 标题
        if line.startswith('# '):
            _add_h1(doc, _strip_md(line[2:]))
        elif line.startswith('## '):
            _add_h2(doc, _strip_md(line[3:]))
        elif line.startswith('### '):
            _add_h3(doc, _strip_md(line[4:]))
        elif line.startswith('#### '):
            _add_h4(doc, _strip_md(line[5:]))

        # 分割线
        elif re.match(r'^[-*]{3,}\s*$', line):
            _add_divider(doc)

        # 列表
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_rich_text(p, line[2:])

        elif re.match(r'^\d+[.、] ', line):
            text = re.sub(r'^\d+[.、] ', '', line)
            p = doc.add_paragraph(style='List Number')
            _add_rich_text(p, text)

        # 表格
        elif line.startswith('|') and '|' in line[1:]:
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip())
                i += 1
            _add_table(doc, rows)
            continue

        # 普通段落
        else:
            p = doc.add_paragraph()
            _add_rich_text(p, line)

        i += 1

    return doc


# ════════════════════════════════════════════════════════════════
# 封面（VS Code 风格：简洁克制）
# ════════════════════════════════════════════════════════════════

def _add_cover(doc, stock_name: str, trade_date: str, md_text: str):
    """添加简洁封面页"""
    if not stock_name:
        m = re.search(r'^# (.+?)(?:（.*?）)?午间分析报告', md_text, re.M)
        if m:
            stock_name = _strip_md(m.group(1)).strip()
    if not trade_date:
        m = re.search(r'\*+\s*数据截止[：:]\s*(\S+)', md_text)
        if m:
            trade_date = m.group(1)

    # 垂直居中（上方留白）
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # 报告标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('午间分析报告')
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = C_H1
    run.font.name = 'Microsoft YaHei'

    # 股票名 + 分隔线
    if stock_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(stock_name)
        run.font.size = Pt(14)
        run.font.color.rgb = C_MUTED
        run.font.name = 'Microsoft YaHei'

    # 细分隔线（仿 VS Code H1 底部线）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="6" w:color="{C_DIVIDER}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # 日期
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f'{trade_date or datetime.now().strftime("%Y-%m-%d")}')
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_MUTED
    run.font.name = 'Microsoft YaHei'

    # 免责声明
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('本报告由 AI 自动生成 · 仅供参考 · 不构成投资建议')
    run.font.size = Pt(8.5)
    run.font.color.rgb = C_LIGHT
    run.italic = True
    run.font.name = 'Microsoft YaHei'

    doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 页眉页脚（简约）
# ════════════════════════════════════════════════════════════════

def _add_header_footer(doc, stock_name: str):
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'{stock_name or "午间分析"}  |  数据截止 11:30')
    run.font.size = Pt(7.5)
    run.font.color.rgb = C_LIGHT
    run.font.name = 'Microsoft YaHei'

    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— ')
    run.font.size = Pt(7.5)
    run.font.color.rgb = C_LIGHT
    fld = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld)
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run._r.append(instr)
    run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    r2 = p.add_run(' —')
    r2.font.size = Pt(7.5)
    r2.font.color.rgb = C_LIGHT


# ════════════════════════════════════════════════════════════════
# 样式
# ════════════════════════════════════════════════════════════════

def _set_styles(doc):
    """设置文档样式（仿 VS Code Markdown Preview）"""
    # ── Normal（正文） ──
    s = doc.styles['Normal']
    s.font.name = 'Microsoft YaHei'
    s.font.size = Pt(10.5)
    s.font.color.rgb = C_BODY
    s.paragraph_format.line_spacing = 1.5
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.space_before = Pt(0)
    _set_east_asia(s, 'Microsoft YaHei')

    # ── 标题用函数 _add_h1/_add_h2 ...，不依赖 style 预设
    # ── List ──
    for style_name in ['List Bullet', 'List Number']:
        s = doc.styles[style_name]
        s.font.name = 'Microsoft YaHei'
        s.font.size = Pt(10.5)
        s.font.color.rgb = C_BODY
        s.paragraph_format.space_after = Pt(2)
        _set_east_asia(s, 'Microsoft YaHei')


def _set_east_asia(style, font_name: str):
    style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


# ════════════════════════════════════════════════════════════════
# 标题（仿 VS Code Markdown Preview：H1/H2 带底线）
# ════════════════════════════════════════════════════════════════

def _add_h1(doc, text: str):
    """H1：大号加粗 + 底部通栏细线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = C_H1
    run.font.name = 'Microsoft YaHei'
    # 底部细线
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="{C_DIVIDER}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_h2(doc, text: str):
    """H2：同 H1 风格，略小"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = C_H2
    run.font.name = 'Microsoft YaHei'
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="{C_DIVIDER}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_h3(doc, text: str):
    """H3：加粗，无底线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = C_H3
    run.font.name = 'Microsoft YaHei'


def _add_h4(doc, text: str):
    """H4"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = C_BODY
    run.font.name = 'Microsoft YaHei'


# ════════════════════════════════════════════════════════════════
# 分割线
# ════════════════════════════════════════════════════════════════

def _add_divider(doc):
    """水平分割线（居中窄线，仿 GitHub/VS Code）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{C_DIVIDER}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


# ════════════════════════════════════════════════════════════════
# 表格（仿 GitHub Markdown 风格）
# ════════════════════════════════════════════════════════════════

def _add_table(doc, rows: list[str]):
    """添加 Markdown 表格（GitHub 风格）"""
    data_rows = [r for r in rows if not re.match(r'^[\s|:—\-]+$', r)]
    if len(data_rows) < 1:
        return

    parsed = [_parse_table_row(r) for r in data_rows]
    if not parsed:
        return

    num_cols = max(len(row) for row in parsed)
    if num_cols < 2:
        return

    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_idx, cells in enumerate(parsed):
        for col_idx, cell_text in enumerate(cells):
            if col_idx >= num_cols:
                break
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.2

            if row_idx == 0:
                # 表头：浅灰底 + 加粗
                _set_cell_shading(cell, C_TABLE_HEADER)
                run = p.add_run(cell_text.strip())
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = C_H1
                run.font.name = 'Microsoft YaHei'
            else:
                if row_idx % 2 == 0:
                    _set_cell_shading(cell, C_TABLE_ALT)
                _add_rich_text(p, cell_text.strip(), font_size=Pt(9))

    doc.add_paragraph()


def _parse_table_row(row: str) -> list[str]:
    """解析 markdown 表格行"""
    row = row.strip()
    if not row.startswith('|'):
        return []
    row = row[1:]
    if row.endswith('|'):
        row = row[:-1]
    cells = []
    current = ''
    for ch in row:
        if ch == '|':
            cells.append(current.strip())
            current = ''
        else:
            current += ch
    cells.append(current.strip())
    return cells


def _set_cell_shading(cell, color: str):
    """设置单元格底纹"""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


# ════════════════════════════════════════════════════════════════
# 富文本（粗体/普通文字）
# ════════════════════════════════════════════════════════════════

def _add_rich_text(paragraph, text: str, font_size: Pt | None = None):
    """添加带粗体格式的文本（粗体用 C_H1 色突出）"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Microsoft YaHei'
            if font_size:
                run.font.size = font_size
            run.font.color.rgb = C_H1
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Microsoft YaHei'
            if font_size:
                run.font.size = font_size


def _strip_md(text: str) -> str:
    """移除 markdown 标记"""
    return text.replace('**', '').replace('*', '').strip()


# ════════════════════════════════════════════════════════════════
# 入口
# ════════════════════════════════════════════════════════════════

def convert_file(input_path: str, output_path: str | None = None):
    """转换单个 md 文件为 docx"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '.docx'
    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    stock_dir = os.path.basename(os.path.dirname(input_path))
    doc = md_to_docx(md_text, stock_name=stock_dir)
    doc.save(output_path)


def convert_directory(dir_path: str):
    """批量转换"""
    for name in sorted(os.listdir(dir_path)):
        sub = os.path.join(dir_path, name)
        md_file = os.path.join(sub, f"{name}_midday.md") if os.path.isdir(sub) else None
        if md_file and os.path.exists(md_file) and name != 'test':
            convert_file(md_file)


def main():
    parser = argparse.ArgumentParser(description='MD→DOCX（VS Code Markdown 风格）')
    parser.add_argument('input', help='输入文件')
    parser.add_argument('output', nargs='?', help='输出文件（可选）')
    args = parser.parse_args()
    convert_file(args.input, args.output or args.input.replace('.md', '.docx'))


if __name__ == '__main__':
    main()
