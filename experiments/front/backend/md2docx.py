#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
md2docx.py — 可复用 Markdown → Word 转换工具
=============================================
用法:
    python md2docx.py <input.md> [output.docx]   # 单个文件
    python md2docx.py <input_dir>                # 批量转换目录下所有 .md

输出:
    - 单个文件: 未指定 output.docx 时, 输出到输入文件同目录同名 .docx
    - 批量模式: 每个 .md 在同目录生成同名 .docx

支持语法:
    # / ## / ### 标题, **加粗**, `行内代码`, | 表格 |, - 列表, --- 分隔线

依赖:
    pip install python-docx   (conda 环境 stock_agent 已预装)

示例:
    python md2docx.py 绿的谐波mx公司分析报告.md
    python md2docx.py 绿的谐波mx公司分析报告.md 输出.docx
    python md2docx.py /home/stockagent/project_space/demand/cases
"""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

BODY_FONT = "微软雅黑"          # 正文字体 (含中文)
BODY_SIZE = 10.5                # 正文字号 (pt)
HEADING_SIZES = {1: 16, 2: 13, 3: 11.5}
HEADING_COLOR = (0x1F, 0x3B, 0x63)   # 深蓝
TABLE_HEADER_FILL = "D9E2F3"         # 表头底色 (浅蓝)
TABLE_FONT_SIZE = 9.5


def set_run_font(run, size=BODY_SIZE, bold=False, color=None, name=BODY_FONT):
    """统一设置 run 字体（含中文字体 eastAsia）。"""
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_markdown_runs(paragraph, text, base_bold=False, size=BODY_SIZE):
    """解析行内 **加粗** 与 `代码` 标记并写入段落。"""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            set_run_font(paragraph.add_run(part[2:-2]), size=size, bold=True)
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            set_run_font(paragraph.add_run(part[1:-1]), size=size)
        else:
            set_run_font(paragraph.add_run(part), size=size, bold=base_bold)


def add_heading(doc, text, level):
    h = doc.add_heading(level=min(level, 3))
    for run in h.runs:
        run.font.color.rgb = RGBColor(*HEADING_COLOR)
        set_run_font(run, size=HEADING_SIZES.get(level, 11.5), bold=True)
    # 标题文本若含 markdown 标记（罕见），补一次解析
    if '**' in text or '`' in text:
        h.text = ''
        add_markdown_runs(h, text, size=HEADING_SIZES.get(level, 11.5))
    return h


def add_table(doc, rows):
    """rows: list[list[str]]，首行为表头。"""
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            add_markdown_runs(p, row[j] if j < len(row) else '',
                              base_bold=(i == 0), size=TABLE_FONT_SIZE)
            if i == 0:
                shd = cell._element.get_or_add_tcPr().makeelement(
                    qn('w:shd'), {qn('w:val'): 'clear',
                                  qn('w:fill'): TABLE_HEADER_FILL})
                cell._element.get_or_add_tcPr().append(shd)
    return table


def convert_md_to_docx(md_path, docx_path):
    """转换单个 md 文件为 docx。"""
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_SIZE)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if re.match(r'^-{3,}$', stripped):          # --- 分隔线
            i += 1
            continue
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)  # 标题
        if m:
            add_heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1
            continue
        if stripped.startswith('|'):                # 表格块
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip('|').split('|')]
                if all(re.match(r'^:?-{2,}:?$', c) for c in cells):
                    continue                        # 跳过对齐分隔行
                rows.append(cells)
            if rows:
                add_table(doc, rows)
            continue
        if re.match(r'^[-*]\s+', stripped):         # 无序列表
            p = doc.add_paragraph(style='List Bullet')
            add_markdown_runs(p, re.sub(r'^[-*]\s+', '', stripped))
            i += 1
            continue
        if not stripped:                            # 空行
            i += 1
            continue
        p = doc.add_paragraph()                     # 普通段落
        add_markdown_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    return docx_path


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    src = sys.argv[1]
    if os.path.isdir(src):
        mds = sorted(f for f in os.listdir(src) if f.lower().endswith('.md'))
        if not mds:
            print(f"目录中未找到 .md 文件: {src}")
            sys.exit(1)
        for name in mds:
            md = os.path.join(src, name)
            out = os.path.join(src, os.path.splitext(name)[0] + '.docx')
            print(f"转换: {name} -> {os.path.basename(out)}")
            convert_md_to_docx(md, out)
        print(f"完成, 共 {len(mds)} 个文件。")
    else:
        md = src
        if not os.path.isfile(md):
            print(f"文件不存在: {md}")
            sys.exit(1)
        out = sys.argv[2] if len(sys.argv) > 2 else \
            os.path.splitext(md)[0] + '.docx'
        convert_md_to_docx(md, out)
        print(f"完成: {out}")


if __name__ == '__main__':
    main()
